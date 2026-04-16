"""Document parsing helpers for the knowledge-base module."""

from __future__ import annotations

import re
from html.parser import HTMLParser
from pathlib import Path
from typing import Any


class _HTMLTextExtractor(HTMLParser):
    """Convert rendered Markdown HTML into plain text."""

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []

    def handle_data(self, data: str) -> None:
        self._parts.append(data)

    def get_text(self) -> str:
        return "".join(self._parts)


def _normalize_text(text: str) -> str:
    """Collapse excessive blank lines while keeping paragraph breaks."""
    cleaned = text.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def extract_text_from_pdf(file_path: str | Path) -> str:
    """Extract plain text from a PDF file."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - depends on optional package
        raise RuntimeError("缺少 pypdf 依赖，无法解析 PDF。") from exc

    reader = PdfReader(str(file_path))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return _normalize_text("\n\n".join(filter(None, pages)))


def extract_text_from_docx(file_path: str | Path) -> str:
    """Extract plain text from a Word document."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - depends on optional package
        raise RuntimeError("缺少 python-docx 依赖，无法解析 DOCX。") from exc

    document = Document(str(file_path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return _normalize_text("\n\n".join(paragraphs))


def extract_text_from_markdown(file_path: str | Path) -> str:
    """Extract readable text from a Markdown file."""
    raw_text = Path(file_path).read_text(encoding="utf-8", errors="ignore")

    try:
        import markdown as markdown_lib
    except ImportError:  # pragma: no cover - depends on optional package
        return _normalize_text(raw_text)

    extractor = _HTMLTextExtractor()
    extractor.feed(markdown_lib.markdown(raw_text))
    text = extractor.get_text().strip()
    return _normalize_text(text or raw_text)


def split_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[dict[str, Any]]:
    """Split text into overlapping chunks for embedding."""
    normalized = _normalize_text(text)
    if not normalized:
        return []
    if chunk_size <= 0:
        raise ValueError("chunk_size 必须大于 0。")
    if overlap < 0 or overlap >= chunk_size:
        raise ValueError("overlap 必须在 0 到 chunk_size 之间。")

    chunks: list[dict[str, Any]] = []
    start = 0
    index = 0
    text_length = len(normalized)

    while start < text_length:
        end = min(start + chunk_size, text_length)
        snippet = normalized[start:end].strip()
        if snippet:
            chunks.append({"text": snippet, "index": index})
            index += 1
        if end >= text_length:
            break
        start = max(end - overlap, start + 1)

    return chunks


def parse_document(file_path: str | Path) -> tuple[str, dict[str, Any]]:
    """Extract text and lightweight metadata from a supported document."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - depends on optional package
            raise RuntimeError("缺少 pypdf 依赖，无法解析 PDF。") from exc

        reader = PdfReader(str(path))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        text = _normalize_text("\n\n".join(filter(None, pages)))
        return text, {"page_count": len(reader.pages), "parser": "pdf"}

    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - depends on optional package
            raise RuntimeError("缺少 python-docx 依赖，无法解析 DOCX。") from exc

        document = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        text = _normalize_text("\n\n".join(paragraphs))
        return text, {"paragraph_count": len(paragraphs), "parser": "docx"}

    if suffix in {".md", ".markdown"}:
        text = extract_text_from_markdown(path)
        paragraph_count = len([item for item in text.split("\n\n") if item.strip()])
        return text, {"paragraph_count": paragraph_count, "parser": "markdown"}

    raise ValueError("仅支持 PDF、DOCX、Markdown 文档。")
