"""Lesson-plan export service for Word and PDF outputs."""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from ..models import Plan
from ..schemas import MiniGamePayload
from .plan_service import PlanService

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover - exercised in environments without python-docx.
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    REPORTLAB_AVAILABLE = True
except ImportError:  # pragma: no cover - exercised in environments without reportlab.
    REPORTLAB_AVAILABLE = False


IMAGE_PLACEHOLDER_PATTERN = re.compile(r"!\[(?P<label>[^\]]*)\]\((?P<target>[^)]*)\)")
logger = logging.getLogger(__name__)


class ExportError(Exception):
    """Base class for export failures."""


class PlanNotFoundError(ExportError):
    """Raised when the requested lesson plan does not exist."""


class ExportUnavailableError(ExportError):
    """Raised when the requested export format is unavailable."""


@dataclass(slots=True)
class ContentBlock:
    """Normalized content block ready for rendering."""

    kind: str
    text: str


@dataclass(slots=True)
class SectionContent:
    """Normalized lesson-plan section."""

    title: str
    blocks: list[ContentBlock]


class ExportService:
    """Generate lesson plans as downloadable files."""

    DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    PDF_MEDIA_TYPE = "application/pdf"
    DEFAULT_TEMPLATE = "default"

    def __init__(self, plan_service: PlanService):
        self.plan_service = plan_service

    def export_to_docx(self, plan_id: str, template: str = DEFAULT_TEMPLATE) -> bytes:
        """Generate a Word document for the given lesson plan."""
        plan = self._get_plan_or_raise(plan_id)
        return self.render_plan_to_docx(plan, template=template)

    def render_plan_to_docx(self, plan: Plan, template: str = DEFAULT_TEMPLATE) -> bytes:
        """Generate a Word document from an already loaded plan."""
        self._ensure_supported_template(template)
        if not DOCX_AVAILABLE:
            raise ExportUnavailableError("Word 导出依赖未安装，请安装 python-docx。")

        document = Document()
        self._configure_docx_document(document)
        sections = self._extract_sections(plan)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(plan.title)
        title_run.bold = True
        title_run.font.size = Pt(18)
        self._set_docx_font(title_run, "SimSun")

        subtitle_bits = [part for part in (plan.subject, plan.grade) if part]
        if subtitle_bits:
            subtitle = document.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.add_run(" / ".join(subtitle_bits))
            subtitle_run.italic = True
            subtitle_run.font.size = Pt(10.5)
            self._set_docx_font(subtitle_run, "SimSun")

        for section in sections:
            document.add_heading(section.title, level=2)
            for block in section.blocks:
                paragraph = document.add_paragraph(
                    style="List Bullet" if block.kind == "bullet" else "List Number" if block.kind == "number" else None
                )
                run = paragraph.add_run(block.text)
                self._set_docx_font(run, "SimSun")

        if not sections:
            paragraph = document.add_paragraph("暂无教案内容。")
            self._set_docx_font(paragraph.runs[0], "SimSun")

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def export_to_pdf(self, plan_id: str, template: str = DEFAULT_TEMPLATE) -> bytes:
        """Generate a PDF document for the given lesson plan."""
        plan = self._get_plan_or_raise(plan_id)
        return self.render_plan_to_pdf(plan, template=template)

    def render_plan_to_pdf(self, plan: Plan, template: str = DEFAULT_TEMPLATE) -> bytes:
        """Generate a PDF document from an already loaded plan."""
        self._ensure_supported_template(template)
        if not REPORTLAB_AVAILABLE:
            raise ExportUnavailableError("PDF 导出依赖未安装，请在 lesson_agent_env 中安装 reportlab。")

        font_name = self._register_pdf_font()
        stylesheet = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "LessonTitle",
            parent=stylesheet["Heading1"],
            fontName=font_name,
            fontSize=18,
            leading=24,
            alignment=TA_CENTER,
            spaceAfter=10,
        )
        meta_style = ParagraphStyle(
            "LessonMeta",
            parent=stylesheet["Normal"],
            fontName=font_name,
            fontSize=10,
            leading=14,
            alignment=TA_CENTER,
            spaceAfter=16,
        )
        heading_style = ParagraphStyle(
            "LessonSectionHeading",
            parent=stylesheet["Heading2"],
            fontName=font_name,
            fontSize=14,
            leading=18,
            spaceBefore=6,
            spaceAfter=8,
        )
        body_style = ParagraphStyle(
            "LessonBody",
            parent=stylesheet["Normal"],
            fontName=font_name,
            fontSize=11,
            leading=17,
            spaceAfter=6,
        )
        bullet_style = ParagraphStyle(
            "LessonBullet",
            parent=body_style,
            leftIndent=12,
            firstLineIndent=0,
        )

        story: list[Any] = [Paragraph(self._escape_text(plan.title), title_style)]
        subtitle_bits = [part for part in (plan.subject, plan.grade) if part]
        if subtitle_bits:
            story.append(Paragraph(self._escape_text(" / ".join(subtitle_bits)), meta_style))

        sections = self._extract_sections(plan)
        if not sections:
            story.append(Paragraph("暂无教案内容。", body_style))
        else:
            for section in sections:
                story.append(Paragraph(self._escape_text(section.title), heading_style))
                list_items: list[ListItem] = []
                current_list_kind: str | None = None
                for block in section.blocks:
                    if block.kind in {"bullet", "number"}:
                        if current_list_kind is not None and current_list_kind != block.kind and list_items:
                            story.append(self._build_pdf_list(list_items, current_list_kind))
                            list_items = []
                        current_list_kind = block.kind
                        list_items.append(ListItem(Paragraph(self._escape_text(block.text), bullet_style)))
                        continue

                    if list_items:
                        story.append(self._build_pdf_list(list_items, current_list_kind or "bullet"))
                        list_items = []
                        current_list_kind = None

                    story.append(Paragraph(self._escape_text(block.text), body_style))

                if list_items:
                    story.append(self._build_pdf_list(list_items, current_list_kind or "bullet"))
                story.append(Spacer(1, 6))

        buffer = BytesIO()
        document = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=42,
            rightMargin=42,
            topMargin=48,
            bottomMargin=48,
            title=plan.title,
        )
        document.build(story)
        return buffer.getvalue()

    def _get_plan_or_raise(self, plan_id: str) -> Plan:
        plan = self.plan_service.get(plan_id)
        if plan is None:
            raise PlanNotFoundError("教案不存在。")
        return plan

    def _extract_sections(self, plan: Plan) -> list[SectionContent]:
        content = plan.content or {}
        raw_sections = content.get("sections", []) if isinstance(content, dict) else []
        sections: list[SectionContent] = []

        for index, raw_section in enumerate(raw_sections, start=1):
            if isinstance(raw_section, dict):
                title = self._coerce_text(
                    raw_section.get("title")
                    or raw_section.get("type")
                    or raw_section.get("name")
                    or raw_section.get("heading")
                    or f"章节 {index}"
                )
                source = raw_section.get("content")
                if source is None:
                    source = raw_section.get("items")
            else:
                title = f"章节 {index}"
                source = raw_section

            blocks = self._normalize_content(source)
            sections.append(SectionContent(title=title, blocks=blocks or [ContentBlock(kind="paragraph", text="暂无内容。")]))

        sections.extend(self._extract_game_sections(content))
        return sections

    def _extract_game_sections(self, content: dict[str, Any]) -> list[SectionContent]:
        """Append generated mini-games as exportable lesson sections."""
        raw_games = content.get("games")
        if not isinstance(raw_games, list) or not raw_games:
            return []

        blocks: list[ContentBlock] = []
        for index, item in enumerate(raw_games, start=1):
            try:
                game = MiniGamePayload.model_validate(item)
            except Exception:
                continue

            blocks.append(ContentBlock(kind="paragraph", text=f"{index}. {game.title or f'小游戏 {index}'}"))
            if game.description:
                blocks.append(ContentBlock(kind="paragraph", text=f"玩法：{game.description}"))
            if game.learning_goal:
                blocks.append(ContentBlock(kind="paragraph", text=f"知识点：{game.learning_goal}"))
            if game.source_section:
                blocks.append(ContentBlock(kind="paragraph", text=f"来源章节：{game.source_section}"))
            for line in self._summarize_game_data(game):
                blocks.append(ContentBlock(kind="bullet", text=line))
            if game.html_url:
                blocks.append(ContentBlock(kind="paragraph", text=f"互动页面：{game.html_url}"))

        return [SectionContent(title="课堂小游戏", blocks=blocks)] if blocks else []

    def _summarize_game_data(self, game: MiniGamePayload) -> list[str]:
        """Convert structured game data into export-friendly bullets."""
        if game.template == "single_choice":
            questions = game.data.get("questions")
            if isinstance(questions, list):
                return [
                    f"{index}. {self._coerce_text(item.get('stem'))}"
                    for index, item in enumerate(questions, start=1)
                    if isinstance(item, dict) and self._coerce_text(item.get("stem"))
                ]
        if game.template == "true_false":
            statements = game.data.get("statements")
            if isinstance(statements, list):
                return [
                    f"{index}. {self._coerce_text(item.get('statement'))}"
                    for index, item in enumerate(statements, start=1)
                    if isinstance(item, dict) and self._coerce_text(item.get("statement"))
                ]
        cards = game.data.get("cards")
        if isinstance(cards, list):
            return [
                f"{self._coerce_text(item.get('front'))}：{self._coerce_text(item.get('back'))}"
                for item in cards
                if isinstance(item, dict)
                and self._coerce_text(item.get("front"))
                and self._coerce_text(item.get("back"))
            ]
        return []

    def _normalize_content(self, value: Any) -> list[ContentBlock]:
        if value is None:
            return []
        if isinstance(value, str):
            blocks: list[ContentBlock] = []
            for line in value.splitlines():
                parsed = self._parse_text_line(line)
                if parsed is not None:
                    blocks.append(parsed)
            return blocks
        if isinstance(value, list):
            blocks: list[ContentBlock] = []
            for item in value:
                if isinstance(item, dict):
                    item_title = self._coerce_text(item.get("title") or item.get("label") or item.get("name"))
                    item_content = self._coerce_text(
                        item.get("content") or item.get("text") or item.get("description") or item.get("value")
                    )
                    combined = "：".join(part for part in (item_title, item_content) if part)
                    if combined:
                        blocks.append(ContentBlock(kind="bullet", text=combined))
                    continue

                item_text = self._coerce_text(item)
                if item_text:
                    blocks.append(ContentBlock(kind="bullet", text=self._replace_image_placeholders(item_text)))
            return blocks
        if isinstance(value, dict):
            blocks: list[ContentBlock] = []
            for key, item in value.items():
                item_text = self._coerce_text(item)
                if item_text:
                    label = self._coerce_text(key)
                    combined = f"{label}：{item_text}" if label else item_text
                    blocks.append(ContentBlock(kind="paragraph", text=self._replace_image_placeholders(combined)))
            return blocks

        item_text = self._coerce_text(value)
        return [ContentBlock(kind="paragraph", text=self._replace_image_placeholders(item_text))] if item_text else []

    def _parse_text_line(self, line: str) -> ContentBlock | None:
        text = line.strip()
        if not text:
            return None

        text = self._replace_image_placeholders(text)
        if re.match(r"^[-*•]\s+", text):
            return ContentBlock(kind="bullet", text=re.sub(r"^[-*•]\s+", "", text, count=1))
        if re.match(r"^\d+[.)]\s+", text):
            return ContentBlock(kind="number", text=re.sub(r"^\d+[.)]\s+", "", text, count=1))
        return ContentBlock(kind="paragraph", text=text)

    def _replace_image_placeholders(self, text: str) -> str:
        def repl(match: re.Match[str]) -> str:
            label = (match.group("label") or "").strip()
            target = (match.group("target") or "").strip()
            if label and target and target != "upload_needed":
                return f"[图片占位：{label}，来源 {target}]"
            if label:
                return f"[图片占位：{label}]"
            return "[图片占位]"

        return IMAGE_PLACEHOLDER_PATTERN.sub(repl, text)

    def _configure_docx_document(self, document: Any) -> None:
        section = document.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "SimSun"
        normal_style.font.size = Pt(11)
        style_properties = normal_style._element.get_or_add_rPr()
        style_properties.rFonts.set(qn("w:eastAsia"), "SimSun")

    def _set_docx_font(self, run: Any, font_name: str) -> None:
        run.font.name = font_name
        run_properties = run._element.get_or_add_rPr()
        run_properties.rFonts.set(qn("w:eastAsia"), font_name)

    def _ensure_supported_template(self, template: str) -> None:
        if template != self.DEFAULT_TEMPLATE:
            raise ExportUnavailableError(f"暂不支持模板 {template!r}。")

    def _register_pdf_font(self) -> str:
        if not REPORTLAB_AVAILABLE:
            raise ExportUnavailableError("PDF 导出不可用。")

        preferred_paths = [
            Path("/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf"),
            Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
            Path("/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc"),
            Path("/usr/share/fonts/truetype/msyh.ttc"),
            Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
            Path("/usr/share/fonts/truetype/arphic/uming.ttc"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/System/Library/Fonts/PingFang.ttc"),
            Path("/Library/Fonts/Arial Unicode.ttf"),
            Path("C:/Windows/Fonts/msyh.ttc"),
            Path("C:/Windows/Fonts/simhei.ttf"),
            Path("C:/Windows/Fonts/simsun.ttc"),
        ]

        font_name = "LessonPlanFont"
        for path in preferred_paths:
            if not path.exists():
                continue
            try:
                pdfmetrics.registerFont(TTFont(font_name, str(path)))
                logger.info("PDF export font registered: %s", path)
                return font_name
            except Exception as exc:  # noqa: BLE001
                logger.warning("Failed to register PDF font %s: %s", path, exc)
                continue

        try:
            fallback_font = "STSong-Light"
            pdfmetrics.registerFont(UnicodeCIDFont(fallback_font))
            logger.warning("PDF export fell back to built-in CJK font %s.", fallback_font)
            return fallback_font
        except Exception as exc:  # noqa: BLE001
            logger.warning("Failed to register built-in CJK PDF font: %s", exc)

        logger.warning("PDF export fell back to Helvetica; Chinese glyphs may not render correctly.")
        return "Helvetica"

    def _build_pdf_list(self, items: list[Any], list_kind: str) -> Any:
        bullet_type = "1" if list_kind == "number" else "bullet"
        return ListFlowable(items, bulletType=bullet_type, leftIndent=18)

    def _coerce_text(self, value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, str):
            return value.strip()
        return str(value).strip()

    def _escape_text(self, text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )
